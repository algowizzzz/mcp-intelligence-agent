"""
Operational Tools Suite — pdf_read, md_save, md_to_docx, search_files, fill_template, list_versions
"""
import io
import os, re, json, shutil, mimetypes
from datetime import datetime, timezone
from pathlib import Path
from sajha.tools.base_mcp_tool import BaseMCPTool
from sajha.core.properties_configurator import PropertiesConfigurator
from sajha.storage import storage
from sajha.path_resolver import resolve as path_resolve


def _get_worker_ctx():
    try:
        from flask import g as _g
        return getattr(_g, 'worker_ctx', {}) or {}
    except RuntimeError:
        return {}


def _props():
    return PropertiesConfigurator()

def _resolve(path_str):
    """Resolve a path relative to the sajhamcpserver working directory (CWD)."""
    p = Path(path_str)
    if p.is_absolute():
        return p.resolve()
    return (Path.cwd() / p).resolve()

def _domain_root():
    """Return domain data root. Checks per-request worker context first (REQ-API-02)."""
    try:
        from flask import g as _g
        worker_root = getattr(_g, 'worker_data_root', None)
        if worker_root:
            return Path(worker_root.rstrip('/')).resolve()
    except RuntimeError:
        pass
    return _resolve(_props().get('data.domain_data.dir', './data/domain_data'))

def _my_data_root():
    """Return my-data root. Checks per-request worker context first (REQ-DD-02 / REQ-MD-01)."""
    try:
        from flask import g as _g
        my_data_root = getattr(_g, 'worker_my_data_root', None)
        if my_data_root:
            return Path(my_data_root.rstrip('/')).resolve()
    except RuntimeError:
        pass
    return _resolve(_props().get('data.my_data.dir', './data/uploads'))

def _templates_dir():
    """Return templates dir. Checks per-request worker context first (REQ-API-02)."""
    try:
        from flask import g as _g
        worker_root = getattr(_g, 'worker_data_root', None)
        if worker_root:
            return Path(worker_root.rstrip('/') + '/templates').resolve()
    except RuntimeError:
        pass
    return _resolve(_props().get('data.templates_dir', './data/domain_data/templates'))

def _safe_path(path_str, *allowed_roots):
    """Return resolved Path if it's within one of the allowed roots, else None."""
    try:
        p = Path(path_str).resolve()
    except Exception:
        return None
    for root in allowed_roots:
        root = Path(root).resolve()
        try:
            p.relative_to(root)
            return p
        except ValueError:
            pass
    return None


# ─────────────────────────── pdf_read ────────────────────────────────────────

class PdfReadTool(BaseMCPTool):
    def __init__(self, config=None):
        cfg = {
            "name": "pdf_read",
            "description": (
                "Extract text and tables from a local PDF file. "
                "Pass the absolute file_path from list_data_files or list_uploaded_files. "
                "Use pages='1-5' for large PDFs to limit output size."
            ),
            "version": "1.0.0",
            "enabled": True,
        }
        if config:
            cfg.update(config)
        super().__init__(cfg)

    def get_input_schema(self):
        return {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Absolute path to a .pdf file."},
                "pages": {"type": "string", "description": "Page range: '1', '1-5', or 'all' (default)."},
                "extract_tables": {"type": "boolean", "description": "Extract tables. Default: true."},
                "max_chars": {"type": "integer", "description": "Truncate text at this char count. Default: 50000."},
            },
            "required": ["file_path"],
        }

    def get_output_schema(self):
        return {"type": "object"}

    def execute(self, arguments):
        file_path = arguments.get("file_path", "")
        pages_arg = arguments.get("pages", "all")
        extract_tables = arguments.get("extract_tables", True)
        max_chars = arguments.get("max_chars", 50000)

        safe = _safe_path(file_path, _domain_root(), _my_data_root())
        if not safe:
            return {"error": f"File not found or access denied: {file_path}"}
        if not safe.exists():
            return {"error": f"File not found or access denied: {file_path}"}
        if safe.suffix.lower() != ".pdf":
            return {"error": "pdf_read only accepts .pdf files."}

        try:
            import fitz  # PyMuPDF
        except ImportError:
            try:
                import pdfplumber as _pl
                return self._read_pdfplumber(safe, pages_arg, max_chars)
            except ImportError:
                return {"error": "Neither PyMuPDF nor pdfplumber is installed."}

        try:
            raw = storage.read_bytes(str(safe))
            doc = fitz.open(stream=io.BytesIO(raw), filetype='pdf')
        except Exception as e:
            return {"error": f"PDF could not be parsed: {e}"}

        if doc.is_encrypted:
            return {"error": "PDF is encrypted. Cannot extract without password."}

        total_pages = len(doc)
        # Parse page range
        page_indices = self._parse_pages(pages_arg, total_pages)

        text_parts = []
        tables = []
        for i in page_indices:
            page = doc[i]
            text_parts.append(page.get_text())
            if extract_tables:
                for tab in page.find_tables():
                    headers = tab.header.names if hasattr(tab, 'header') else []
                    rows = tab.extract()
                    if rows:
                        if not headers and rows:
                            headers = rows[0]
                            rows = rows[1:]
                        tables.append({"page": i + 1, "headers": headers, "rows": rows})

        full_text = "\n".join(text_parts)
        truncated = len(full_text) > max_chars
        if truncated:
            full_text = full_text[:max_chars]

        warning = None
        if not full_text.strip():
            warning = "No text layer detected. PDF may be image-only."

        result = {
            "filename": safe.name,
            "pages_extracted": len(page_indices),
            "total_pages": total_pages,
            "char_count": len(full_text),
            "truncated": truncated,
            "text": full_text,
            "_source": str(safe),
        }
        if extract_tables:
            result["tables"] = tables
        if warning:
            result["warning"] = warning
        return result

    def _parse_pages(self, pages_arg, total):
        if pages_arg == "all" or not pages_arg:
            return list(range(total))
        m = re.match(r"^(\d+)-(\d+)$", pages_arg.strip())
        if m:
            a, b = int(m.group(1)) - 1, int(m.group(2)) - 1
            return list(range(max(0, a), min(total, b + 1)))
        m2 = re.match(r"^(\d+)$", pages_arg.strip())
        if m2:
            i = int(m2.group(1)) - 1
            return [i] if 0 <= i < total else []
        return list(range(total))


# ─────────────────────────── md_save ─────────────────────────────────────────

class MdSaveTool(BaseMCPTool):
    def __init__(self, config=None):
        cfg = {
            "name": "md_save",
            "description": (
                "Save a Markdown string to my_data/ as a .md file with automatic versioning. "
                "If the file already exists, the old version is archived with a timestamp suffix. "
                "Use this to persist analysis outputs, filled templates, and canvas content."
            ),
            "version": "1.0.0",
            "enabled": True,
        }
        if config:
            cfg.update(config)
        super().__init__(cfg)

    def get_input_schema(self):
        return {
            "type": "object",
            "properties": {
                "content": {"type": "string", "description": "Markdown content to write."},
                "filename": {"type": "string", "description": "Target filename including .md extension."},
                "subfolder": {"type": "string", "description": "Sub-folder within my_data/. Created if needed. Default: ''."},
                "versioning": {"type": "boolean", "description": "Archive existing file before writing. Default: true."},
                "overwrite": {"type": "boolean", "description": "Overwrite without archiving. Overrides versioning. Default: false."},
            },
            "required": ["content", "filename"],
        }

    def get_output_schema(self):
        return {"type": "object"}

    def execute(self, arguments):
        content = arguments.get("content", "")
        filename = Path(arguments.get("filename", "output.md")).name
        if not filename.endswith(".md"):
            filename += ".md"
        subfolder = arguments.get("subfolder", "")
        versioning = arguments.get("versioning", True)
        overwrite = arguments.get("overwrite", False)

        root = _my_data_root()
        if subfolder:
            folder = root / subfolder
        else:
            folder = root
        folder.mkdir(parents=True, exist_ok=True)

        dest = folder / filename
        archived_as = None

        if dest.exists() and not overwrite:
            if versioning:
                mtime = datetime.fromtimestamp(dest.stat().st_mtime, tz=timezone.utc)
                ts = mtime.strftime("%Y-%m-%d_%H%M%S")
                stem = dest.stem
                archive_name = f"{stem}_{ts}.md"
                archive_path = folder / archive_name
                shutil.move(str(dest), str(archive_path))
                archived_as = archive_name

        storage.write_text(str(dest), content, encoding="utf-8")
        stat = dest.stat()

        return {
            "path": str(dest),
            "filename": filename,
            "subfolder": subfolder,
            "size_bytes": stat.st_size,
            "versioned": archived_as is not None,
            "archived_as": archived_as,
            "written_at": datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
            "_source": str(dest),
        }


# ─────────────────────────── md_to_docx ──────────────────────────────────────

class MdToDocxTool(BaseMCPTool):
    def __init__(self, config=None):
        cfg = {
            "name": "md_to_docx",
            "description": (
                "Convert a .md file to a formatted .docx Word document. "
                "Handles headings, tables, bullet lists, bold/italic, code blocks, and YAML frontmatter (omitted). "
                "Output is saved alongside the source file or to output_path."
            ),
            "version": "1.0.0",
            "enabled": True,
        }
        if config:
            cfg.update(config)
        super().__init__(cfg)

    def get_input_schema(self):
        return {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Absolute path to a .md file."},
                "output_path": {"type": "string", "description": "Where to save .docx. Defaults to alongside source."},
                "style": {"type": "string", "enum": ["standard", "minimal", "report"], "description": "Style preset. Default: standard."},
                "include_toc": {"type": "boolean", "description": "Insert Table of Contents. Default: false."},
            },
            "required": ["file_path"],
        }

    def get_output_schema(self):
        return {"type": "object"}

    def execute(self, arguments):
        file_path = arguments.get("file_path", "")
        output_path = arguments.get("output_path")
        include_toc = arguments.get("include_toc", False)

        safe = _safe_path(file_path, _domain_root(), _my_data_root())
        if not safe or not safe.exists():
            return {"error": f"File not found or access denied: {file_path}"}
        if safe.suffix.lower() != ".md":
            return {"error": "md_to_docx only accepts .md files."}

        text = storage.read_text(str(safe), encoding="utf-8")
        # Strip YAML frontmatter
        text = re.sub(r"^---\n.*?\n---\n", "", text, flags=re.DOTALL)

        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            return {"error": "python-docx is not installed."}

        doc = Document()
        # Set narrow margins
        for section in doc.sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        def set_heading_style(para, level):
            colors = {1: RGBColor(0x1e, 0x3a, 0x5f), 2: RGBColor(0x1e, 0x3a, 0x5f), 3: RGBColor(0x1e, 0x6e, 0xb0)}
            para.style = f'Heading {level}'
            for run in para.runs:
                run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))

        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Headings
            m = re.match(r"^(#{1,6})\s+(.*)", line)
            if m:
                level = min(len(m.group(1)), 3)
                para = doc.add_heading(m.group(2).strip(), level=level)
                i += 1
                continue

            # HR
            if re.match(r"^---+$", line.strip()):
                para = doc.add_paragraph()
                i += 1
                continue

            # Code block
            if line.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # skip closing ```
                para = doc.add_paragraph("\n".join(code_lines))
                para.style = "Normal"
                for run in para.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                continue

            # Table (pipe syntax)
            if "|" in line and line.strip().startswith("|"):
                table_rows = []
                while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                    row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    # Skip separator row
                    if not all(re.match(r"^[-:]+$", c) for c in row if c):
                        table_rows.append(row)
                    i += 1
                if table_rows:
                    max_cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=max_cols)
                    table.style = "Table Grid"
                    for ri, row in enumerate(table_rows):
                        for ci, cell_text in enumerate(row):
                            if ci < max_cols:
                                cell = table.rows[ri].cells[ci]
                                cell.text = cell_text
                                if ri == 0:
                                    for run in cell.paragraphs[0].runs:
                                        run.font.bold = True
                continue

            # Bullet list
            if re.match(r"^[\-\*\+]\s+", line):
                text_part = re.sub(r"^[\-\*\+]\s+", "", line)
                para = doc.add_paragraph(style="List Bullet")
                self._add_inline(para, text_part)
                i += 1
                continue

            # Numbered list
            if re.match(r"^\d+\.\s+", line):
                text_part = re.sub(r"^\d+\.\s+", "", line)
                para = doc.add_paragraph(style="List Number")
                self._add_inline(para, text_part)
                i += 1
                continue

            # Blockquote
            if line.startswith("> "):
                para = doc.add_paragraph(line[2:])
                para.paragraph_format.left_indent = Inches(0.5)
                i += 1
                continue

            # Empty line
            if not line.strip():
                i += 1
                continue

            # Normal paragraph
            para = doc.add_paragraph()
            self._add_inline(para, line)
            i += 1

        if output_path:
            out = Path(output_path)
        else:
            out = safe.with_suffix(".docx")

        doc.save(str(out))
        return {
            "source_md": str(safe),
            "output_docx": str(out),
            "size_bytes": out.stat().st_size,
            "style": arguments.get("style", "standard"),
            "toc_included": include_toc,
            "_source": str(out),
        }

    def _add_inline(self, para, text):
        """Add text with **bold** and *italic* inline formatting."""
        # Split on bold/italic markers
        pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
        parts = re.split(pattern, text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = para.add_run(part[1:-1])
                run.font.name = "Courier New"
                from docx.shared import Pt
                run.font.size = Pt(10)
            else:
                para.add_run(part)


# ─────────────────────────── search_files ────────────────────────────────────

class SearchFilesTool(BaseMCPTool):
    def __init__(self, config=None):
        cfg = {
            "name": "search_files",
            "description": (
                "Full-text search across domain_data/ and my_data/ files. "
                "Returns file path and excerpts for each match. "
                "Supports pdf, docx, xlsx, csv, md, txt, json. "
                "Use to find documents containing a keyword or phrase before processing them."
            ),
            "version": "1.0.0",
            "enabled": True,
        }
        if config:
            cfg.update(config)
        super().__init__(cfg)

    def get_input_schema(self):
        return {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Search term. Case-insensitive. Use quotes for exact phrase."},
                "section": {"type": "string", "enum": ["domain_data", "my_data", "all"], "description": "Limit search scope. Default: all."},
                "file_type": {"type": "string", "description": "Filter by extension: pdf|docx|xlsx|csv|md|txt|json|all. Default: all."},
                "folder": {"type": "string", "description": "Limit to a sub-folder name (partial match)."},
                "max_results": {"type": "integer", "description": "Max file matches. Default: 20."},
                "excerpt_chars": {"type": "integer", "description": "Context chars around each match. Default: 200."},
            },
            "required": ["query"],
        }

    def get_output_schema(self):
        return {"type": "object"}

    def execute(self, arguments):
        query = arguments.get("query", "")
        section = arguments.get("section", "all")
        file_type = arguments.get("file_type", "all")
        folder_filter = arguments.get("folder", "")
        max_results = arguments.get("max_results", 20)
        excerpt_chars = arguments.get("excerpt_chars", 200)

        # Build search pattern
        exact = re.match(r'^"(.+)"$', query)
        if exact:
            pattern = re.compile(re.escape(exact.group(1)), re.IGNORECASE)
        else:
            pattern = re.compile(re.escape(query), re.IGNORECASE)

        # Collect candidate files
        roots = []
        if section in ("domain_data", "all"):
            roots.append(("domain_data", _domain_root()))
        if section in ("my_data", "all"):
            roots.append(("my_data", _my_data_root()))

        candidates = []
        for sec_name, root in roots:
            if not root.exists():
                continue
            for f in root.rglob("*"):
                if not f.is_file():
                    continue
                if f.name.startswith("."):
                    continue
                if folder_filter and folder_filter.lower() not in str(f.parent).lower():
                    continue
                ext = f.suffix.lower().lstrip(".")
                if file_type != "all" and ext != file_type:
                    continue
                if ext in ("parquet", "pq", "db", "wal", "pyc"):
                    continue
                candidates.append((sec_name, f, ext))

        results = []
        for sec_name, f, ext in candidates:
            if len(results) >= max_results:
                break
            try:
                text = self._extract_text(f, ext)
            except Exception:
                continue
            if not text:
                continue
            matches = list(pattern.finditer(text))
            if not matches:
                continue
            excerpts = []
            for m in matches[:5]:
                start = max(0, m.start() - excerpt_chars // 2)
                end = min(len(text), m.end() + excerpt_chars // 2)
                snippet = text[start:end].replace("\n", " ").strip()
                # Highlight match
                snippet = pattern.sub(lambda x: f"[{x.group()}]", snippet)
                excerpts.append(snippet)
            results.append({
                "filename": f.name,
                "path": str(f),
                "file_type": ext,
                "section": sec_name,
                "match_count": len(matches),
                "excerpts": excerpts,
                "_source": str(f),
            })

        return {
            "query": query,
            "total_matches": len(results),
            "results": results,
        }

    def _extract_text(self, path, ext):
        if ext in ("md", "txt", "csv", "tsv"):
            return path.read_text(encoding="utf-8", errors="ignore")
        if ext == "json":
            return path.read_text(encoding="utf-8", errors="ignore")
        if ext == "docx":
            from docx import Document
            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        if ext in ("xlsx", "xls"):
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    parts.append(" ".join(str(c) for c in row if c is not None))
            return "\n".join(parts)
        if ext == "pdf":
            try:
                import fitz
                raw = storage.read_bytes(str(path))
                doc = fitz.open(stream=io.BytesIO(raw), filetype='pdf')
                return "\n".join(doc[i].get_text() for i in range(len(doc)))
            except Exception:
                return ""
        return ""


# ─────────────────────────── fill_template ───────────────────────────────────

class FillTemplateTool(BaseMCPTool):
    def __init__(self, config=None):
        cfg = {
            "name": "fill_template",
            "description": (
                "Read a .md template from domain_data/templates/, substitute {{placeholder}} tokens "
                "with values from a data dict, and save the filled output to my_data/ via md_save. "
                "Templates declare required placeholders in YAML frontmatter. "
                "Use list_data_files with folder='templates' to discover available templates."
            ),
            "version": "1.0.0",
            "enabled": True,
        }
        if config:
            cfg.update(config)
        super().__init__(cfg)

    def get_input_schema(self):
        return {
            "type": "object",
            "properties": {
                "template_path": {"type": "string", "description": "Absolute path to a .md template in domain_data/templates/."},
                "data": {"type": "object", "description": "Key-value dict mapping placeholder names to replacement values."},
                "output_filename": {"type": "string", "description": "Output filename. Defaults to template stem + '_filled_' + timestamp + '.md'."},
                "output_subfolder": {"type": "string", "description": "Sub-folder within my_data/. Default: 'reports'."},
                "versioning": {"type": "boolean", "description": "Archive existing output file. Default: true."},
                "convert_to_docx": {"type": "boolean", "description": "Also convert output to .docx. Default: false."},
            },
            "required": ["template_path", "data"],
        }

    def get_output_schema(self):
        return {"type": "object"}

    def execute(self, arguments):
        template_path = arguments.get("template_path", "")
        data = arguments.get("data", {})
        output_subfolder = arguments.get("output_subfolder", "reports")
        versioning = arguments.get("versioning", True)
        convert_to_docx = arguments.get("convert_to_docx", False)

        # Security: template must be within templates dir
        tmpl_dir = _templates_dir()
        safe = _safe_path(template_path, tmpl_dir)
        if not safe or not safe.exists():
            return {"error": "Template access denied"}
        if safe.suffix.lower() != ".md":
            return {"error": "fill_template only accepts .md template files."}

        content = storage.read_text(str(safe), encoding="utf-8")

        # Strip and parse frontmatter
        fm_match = re.match(r"^---\n(.*?)\n---\n", content, re.DOTALL)
        declared_placeholders = []
        if fm_match:
            fm_text = fm_match.group(1)
            content = content[fm_match.end():]
            # Extract placeholders list from frontmatter
            pm = re.search(r"placeholders:\s*\[([^\]]*)\]", fm_text)
            if pm:
                declared_placeholders = [p.strip().strip("'\"") for p in pm.group(1).split(",") if p.strip()]

        # Substitute placeholders
        filled = content
        for key, value in data.items():
            filled = filled.replace("{{" + key + "}}", str(value))

        # Find missing placeholders (tokens still present)
        remaining = re.findall(r"\{\{(\w+)\}\}", filled)
        missing = list(set(remaining))

        # Determine output filename
        output_filename = arguments.get("output_filename")
        if not output_filename:
            ts = datetime.now(tz=timezone.utc).strftime("%Y%m%d%H%M%S")
            output_filename = f"{safe.stem}_filled_{ts}.md"

        # Save via MdSaveTool logic directly
        saver = MdSaveTool()
        save_result = saver.execute({
            "content": filled,
            "filename": output_filename,
            "subfolder": output_subfolder,
            "versioning": versioning,
        })

        result = {
            "template_used": safe.name,
            "output_md_path": save_result.get("path"),
            "output_docx_path": None,
            "placeholders_filled": len(data) - len(missing),
            "placeholders_missing": missing,
            "versioned": save_result.get("versioned", False),
            "archived_as": save_result.get("archived_as"),
            "_source": save_result.get("path"),
        }

        if convert_to_docx and save_result.get("path"):
            converter = MdToDocxTool()
            docx_result = converter.execute({"file_path": save_result["path"]})
            if "output_docx" in docx_result:
                result["output_docx_path"] = docx_result["output_docx"]

        return result


# ─────────────────────────── list_versions ───────────────────────────────────

class ListVersionsTool(BaseMCPTool):
    def __init__(self, config=None):
        cfg = {
            "name": "list_versions",
            "description": (
                "List all versioned copies of a file stem in my_data/. "
                "Returns canonical (latest) file and all archived versions sorted newest first. "
                "Use after md_save to inspect version history or select a previous version to restore."
            ),
            "version": "1.0.0",
            "enabled": True,
        }
        if config:
            cfg.update(config)
        super().__init__(cfg)

    def get_input_schema(self):
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "Canonical filename (with extension). E.g. 'rbc_ccr_brief.md'."},
                "subfolder": {"type": "string", "description": "Sub-folder within my_data/ to search. Omit to search all of my_data/."},
            },
            "required": ["filename"],
        }

    def get_output_schema(self):
        return {"type": "object"}

    def execute(self, arguments):
        filename = arguments.get("filename", "")
        subfolder = arguments.get("subfolder", "")

        root = _my_data_root()
        if subfolder:
            search_root = root / subfolder
        else:
            search_root = root

        stem = Path(filename).stem
        ext = Path(filename).suffix

        # Pattern for versioned files: stem_YYYY-MM-DD_HHmmss.ext
        version_pattern = re.compile(
            r"^" + re.escape(stem) + r"_(\d{4}-\d{2}-\d{2}_\d{6})" + re.escape(ext) + r"$",
            re.IGNORECASE
        )

        canonical = None
        versions = []

        search_paths = list(search_root.rglob("*")) if search_root.exists() else []
        for f in search_paths:
            if not f.is_file():
                continue
            stat = f.stat()
            mtime = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc)

            if f.name.lower() == filename.lower():
                canonical = {
                    "filename": f.name,
                    "path": str(f),
                    "modified_at": mtime.isoformat(),
                    "size_bytes": stat.st_size,
                }
            else:
                m = version_pattern.match(f.name)
                if m:
                    versions.append({
                        "filename": f.name,
                        "path": str(f),
                        "archived_at": mtime.isoformat(),
                        "size_bytes": stat.st_size,
                        "_ts": m.group(1),
                    })

        # Sort versions newest first
        versions.sort(key=lambda x: x["_ts"], reverse=True)
        for v in versions:
            del v["_ts"]

        if not canonical and versions:
            canonical = dict(versions[0])
            canonical["note"] = "No canonical file found; showing most recent version."

        return {
            "canonical": canonical,
            "versions": versions,
            "version_count": len(versions),
            "_source": str(search_root),
        }
