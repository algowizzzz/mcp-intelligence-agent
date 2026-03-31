"""
Copyright All rights Reserved 2025-2030, Ashutosh Sinha, Email: ajsinha@gmail.com
Microsoft Document Tools MCP Tool Implementation - Refactored with Individual Tools
"""

import os
import json
from typing import Dict, Any, List, Optional
from pathlib import Path
from sajha.tools.base_mcp_tool import BaseMCPTool
from sajha.core.properties_configurator import PropertiesConfigurator


class MsDocBaseTool(BaseMCPTool):
    """
    Base class for Microsoft Document tools with shared functionality
    """
    
    def __init__(self, config: Dict = None):
        """Initialize MS Doc base tool"""
        super().__init__(config)
        
        # Set docs directory from config or use default
        self.docs_directory = (config.get('docs_directory') if config else None) or PropertiesConfigurator().get('tool.msdoc.docs_directory', 'data/msdocs')
        
        # Ensure directory exists
        Path(self.docs_directory).mkdir(parents=True, exist_ok=True)
    
    def _get_file_path(self, filename: str) -> Path:
        """Get full path for a file"""
        return Path(self.docs_directory) / filename
    
    def _list_files_by_type(self, file_type: str = 'all') -> List[Dict]:
        """List files by type"""
        files = []
        
        try:
            for file_path in Path(self.docs_directory).iterdir():
                if file_path.is_file():
                    extension = file_path.suffix.lower()
                    
                    # Filter by type
                    if file_type == 'word' and extension not in ['.docx', '.doc']:
                        continue
                    elif file_type == 'excel' and extension not in ['.xlsx', '.xls', '.xlsm']:
                        continue
                    
                    files.append({
                        'filename': file_path.name,
                        'path': str(file_path),
                        'extension': extension,
                        'size': file_path.stat().st_size,
                        'modified': file_path.stat().st_mtime
                    })
            
            return sorted(files, key=lambda x: x['modified'], reverse=True)
            
        except Exception as e:
            self.logger.error(f"Failed to list files: {e}")
            return []
    
    def _read_word_document(self, file_path: Path) -> Dict:
        """Read Word document"""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract text from tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                'filename': file_path.name,
                'paragraphs': paragraphs,
                'paragraph_count': len(paragraphs),
                'tables': tables,
                'table_count': len(tables)
            }
            
        except ImportError:
            raise ValueError("python-docx library not installed. Install with: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Failed to read Word document: {str(e)}")
    
    def _read_excel_document(self, file_path: Path, sheet_name: Optional[str] = None, 
                            sheet_index: Optional[int] = None, max_rows: int = 100,
                            include_formulas: bool = False) -> Dict:
        """Read Excel document"""
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(str(file_path), data_only=not include_formulas)
            
            # Get sheet
            if sheet_name:
                sheet = wb[sheet_name]
            elif sheet_index is not None:
                sheet = wb.worksheets[sheet_index]
            else:
                sheet = wb.active
            
            # Extract data
            data = []
            for i, row in enumerate(sheet.iter_rows(values_only=not include_formulas)):
                if i >= max_rows:
                    break
                data.append(list(row))
            
            result = {
                'filename': file_path.name,
                'sheet_name': sheet.title,
                'data': data,
                'row_count': len(data),
                'column_count': len(data[0]) if data else 0
            }
            
            if include_formulas:
                formulas = []
                for row in sheet.iter_rows():
                    row_formulas = []
                    for cell in row:
                        if cell.value and isinstance(cell.value, str) and cell.value.startswith('='):
                            row_formulas.append({
                                'cell': cell.coordinate,
                                'formula': cell.value
                            })
                    if row_formulas:
                        formulas.append(row_formulas)
                result['formulas'] = formulas
            
            wb.close()
            return result
            
        except ImportError:
            raise ValueError("openpyxl library not installed. Install with: pip install openpyxl")
        except Exception as e:
            raise ValueError(f"Failed to read Excel document: {str(e)}")


class MsDocListFilesTool(MsDocBaseTool):
    """
    Tool to list available files in the docs directory
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_list_files',
            'description': 'List all Word and Excel documents in the docs directory',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "file_type": {
                    "type": "string",
                    "description": "Type of files to list",
                    "enum": ["all", "word", "excel"],
                    "default": "all"
                }
            }
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "directory": {"type": "string"},
                "file_type": {"type": "string"},
                "count": {"type": "integer"},
                "files": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "filename": {"type": "string"},
                            "path": {"type": "string"},
                            "extension": {"type": "string"},
                            "size": {"type": "integer"},
                            "modified": {"type": "number"}
                        }
                    }
                }
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute list files"""
        file_type = arguments.get('file_type', 'all')
        files = self._list_files_by_type(file_type)
        
        return {
            'directory': self.docs_directory,
            'file_type': file_type,
            'count': len(files),
            'files': files,
            '_source': self.docs_directory
        }


class MsDocReadWordTool(MsDocBaseTool):
    """
    Tool to read Word documents
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_read_word',
            'description': 'Read and extract content from Word documents (.docx)',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Word file to read"
                }
            },
            "required": ["filename"]
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "paragraphs": {
                    "type": "array",
                    "items": {"type": "string"}
                },
                "paragraph_count": {"type": "integer"},
                "tables": {
                    "type": "array",
                    "items": {
                        "type": "array",
                        "items": {
                            "type": "array",
                            "items": {"type": "string"}
                        }
                    }
                },
                "table_count": {"type": "integer"}
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute read Word document"""
        filename = arguments['filename']
        file_path = self._get_file_path(filename)
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")

        result = self._read_word_document(file_path)
        result['_source'] = str(file_path)
        return result


class MsDocReadExcelTool(MsDocBaseTool):
    """
    Tool to read Excel spreadsheets
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_read_excel',
            'description': 'Read and extract data from Excel spreadsheets (.xlsx)',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Excel file to read"
                },
                "sheet_name": {
                    "type": "string",
                    "description": "Name of the sheet to read (optional)"
                },
                "sheet_index": {
                    "type": "integer",
                    "description": "Index of the sheet to read (0-based, optional)",
                    "minimum": 0
                },
                "max_rows": {
                    "type": "integer",
                    "description": "Maximum number of rows to return",
                    "default": 100,
                    "minimum": 1,
                    "maximum": 10000
                },
                "include_formulas": {
                    "type": "boolean",
                    "description": "Include cell formulas",
                    "default": false
                }
            },
            "required": ["filename"]
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "sheet_name": {"type": "string"},
                "data": {
                    "type": "array",
                    "items": {
                        "type": "array"
                    }
                },
                "row_count": {"type": "integer"},
                "column_count": {"type": "integer"},
                "formulas": {
                    "type": "array",
                    "items": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "cell": {"type": "string"},
                                "formula": {"type": "string"}
                            }
                        }
                    }
                }
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute read Excel document"""
        filename = arguments['filename']
        file_path = self._get_file_path(filename)
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")

        result = self._read_excel_document(
            file_path,
            arguments.get('sheet_name'),
            arguments.get('sheet_index'),
            arguments.get('max_rows', 100),
            arguments.get('include_formulas', False)
        )
        result['_source'] = str(file_path)
        return result


class MsDocSearchWordTool(MsDocBaseTool):
    """
    Tool to search for text in Word documents
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_search_word',
            'description': 'Search for text within Word documents',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Word file to search"
                },
                "search_term": {
                    "type": "string",
                    "description": "Text to search for"
                }
            },
            "required": ["filename", "search_term"]
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "search_term": {"type": "string"},
                "matches": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "paragraph_index": {"type": "integer"},
                            "text": {"type": "string"}
                        }
                    }
                },
                "match_count": {"type": "integer"}
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute search in Word document"""
        filename = arguments['filename']
        search_term = arguments['search_term'].lower()
        file_path = self._get_file_path(filename)
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        doc_content = self._read_word_document(file_path)
        
        matches = []
        for i, paragraph in enumerate(doc_content['paragraphs']):
            if search_term in paragraph.lower():
                matches.append({
                    'paragraph_index': i,
                    'text': paragraph
                })
        
        return {
            'filename': filename,
            'search_term': search_term,
            'matches': matches,
            'match_count': len(matches),
            '_source': str(file_path)
        }


class MsDocSearchExcelTool(MsDocBaseTool):
    """
    Tool to search for text in Excel spreadsheets
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_search_excel',
            'description': 'Search for text within Excel spreadsheets',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Excel file to search"
                },
                "search_term": {
                    "type": "string",
                    "description": "Text to search for"
                },
                "sheet_name": {
                    "type": "string",
                    "description": "Name of the sheet to search (optional)"
                }
            },
            "required": ["filename", "search_term"]
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "sheet_name": {"type": "string"},
                "search_term": {"type": "string"},
                "matches": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "row_index": {"type": "integer"},
                            "column_index": {"type": "integer"},
                            "value": {"type": "string"}
                        }
                    }
                },
                "match_count": {"type": "integer"}
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute search in Excel document"""
        filename = arguments['filename']
        search_term = arguments['search_term'].lower()
        sheet_name = arguments.get('sheet_name')
        file_path = self._get_file_path(filename)
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        excel_content = self._read_excel_document(
            file_path,
            sheet_name=sheet_name,
            max_rows=10000
        )
        
        matches = []
        for row_idx, row in enumerate(excel_content['data']):
            for col_idx, cell in enumerate(row):
                if cell and search_term in str(cell).lower():
                    matches.append({
                        'row_index': row_idx,
                        'column_index': col_idx,
                        'value': str(cell)
                    })
        
        return {
            'filename': filename,
            'sheet_name': excel_content['sheet_name'],
            'search_term': search_term,
            'matches': matches,
            'match_count': len(matches),
            '_source': str(file_path)
        }


class MsDocGetWordMetadataTool(MsDocBaseTool):
    """
    Tool to get Word document metadata
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_get_word_metadata',
            'description': 'Get metadata and properties from Word documents',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Word file"
                }
            },
            "required": ["filename"]
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "metadata": {
                    "type": "object",
                    "properties": {
                        "author": {"type": "string"},
                        "title": {"type": "string"},
                        "subject": {"type": "string"},
                        "created": {"type": "string"},
                        "modified": {"type": "string"}
                    }
                }
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute get Word metadata"""
        filename = arguments['filename']
        file_path = self._get_file_path(filename)
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            core_props = doc.core_properties
            
            return {
                'filename': filename,
                'metadata': {
                    'author': core_props.author or '',
                    'title': core_props.title or '',
                    'subject': core_props.subject or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else ''
                },
                '_source': str(file_path)
            }
            
        except ImportError:
            raise ValueError("python-docx library not installed")
        except Exception as e:
            raise ValueError(f"Failed to get metadata: {str(e)}")


class MsDocGetExcelMetadataTool(MsDocBaseTool):
    """
    Tool to get Excel spreadsheet metadata
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_get_excel_metadata',
            'description': 'Get metadata and properties from Excel spreadsheets',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Excel file"
                }
            },
            "required": ["filename"]
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "metadata": {
                    "type": "object",
                    "properties": {
                        "creator": {"type": "string"},
                        "title": {"type": "string"},
                        "subject": {"type": "string"},
                        "created": {"type": "string"},
                        "modified": {"type": "string"}
                    }
                }
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute get Excel metadata"""
        filename = arguments['filename']
        file_path = self._get_file_path(filename)
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(str(file_path))
            props = wb.properties
            
            metadata = {
                'creator': props.creator or '',
                'title': props.title or '',
                'subject': props.subject or '',
                'created': str(props.created) if props.created else '',
                'modified': str(props.modified) if props.modified else ''
            }
            
            wb.close()
            
            return {
                'filename': filename,
                'metadata': metadata,
                '_source': str(file_path)
            }

        except ImportError:
            raise ValueError("openpyxl library not installed")
        except Exception as e:
            raise ValueError(f"Failed to get metadata: {str(e)}")


class MsDocExtractTextTool(MsDocBaseTool):
    """
    Tool to extract plain text from documents
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_extract_text',
            'description': 'Extract all text content from Word or Excel documents',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the file to extract text from"
                }
            },
            "required": ["filename"]
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "text": {"type": "string"},
                "character_count": {"type": "integer"}
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute extract text"""
        filename = arguments['filename']
        file_path = self._get_file_path(filename)
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        extension = file_path.suffix.lower()
        
        if extension in ['.docx', '.doc']:
            doc_content = self._read_word_document(file_path)
            text = '\n'.join(doc_content['paragraphs'])
        elif extension in ['.xlsx', '.xls', '.xlsm']:
            excel_content = self._read_excel_document(file_path, max_rows=10000)
            text_parts = []
            for row in excel_content['data']:
                text_parts.append('\t'.join(str(cell) if cell else '' for cell in row))
            text = '\n'.join(text_parts)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
        
        return {
            'filename': filename,
            'text': text,
            'character_count': len(text),
            '_source': str(file_path)
        }


class MsDocGetExcelSheetsTool(MsDocBaseTool):
    """
    Tool to list all sheets in an Excel workbook
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_get_excel_sheets',
            'description': 'List all sheets in an Excel workbook',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Excel file"
                }
            },
            "required": ["filename"]
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "sheets": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "index": {"type": "integer"},
                            "name": {"type": "string"}
                        }
                    }
                },
                "count": {"type": "integer"}
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute get Excel sheets"""
        filename = arguments['filename']
        file_path = self._get_file_path(filename)
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(str(file_path), read_only=True)
            sheets = [
                {'index': i, 'name': sheet.title}
                for i, sheet in enumerate(wb.worksheets)
            ]
            wb.close()
            
            return {
                'filename': filename,
                'sheets': sheets,
                'count': len(sheets),
                '_source': str(file_path)
            }
            
        except ImportError:
            raise ValueError("openpyxl library not installed")
        except Exception as e:
            raise ValueError(f"Failed to get sheets: {str(e)}")


class MsDocReadExcelSheetTool(MsDocBaseTool):
    """
    Tool to read a specific sheet from an Excel workbook
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_read_excel_sheet',
            'description': 'Read data from a specific sheet in an Excel workbook',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Excel file"
                },
                "sheet_name": {
                    "type": "string",
                    "description": "Name of the sheet to read"
                },
                "sheet_index": {
                    "type": "integer",
                    "description": "Index of the sheet to read (0-based)",
                    "minimum": 0
                },
                "max_rows": {
                    "type": "integer",
                    "description": "Maximum number of rows to return",
                    "default": 100,
                    "minimum": 1,
                    "maximum": 10000
                }
            },
            "required": ["filename"],
            "oneOf": [
                {"required": ["sheet_name"]},
                {"required": ["sheet_index"]}
            ]
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "sheet_name": {"type": "string"},
                "data": {
                    "type": "array",
                    "items": {
                        "type": "array"
                    }
                },
                "row_count": {"type": "integer"},
                "column_count": {"type": "integer"}
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute read Excel sheet"""
        filename = arguments['filename']
        file_path = self._get_file_path(filename)
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")

        result = self._read_excel_document(
            file_path,
            arguments.get('sheet_name'),
            arguments.get('sheet_index'),
            arguments.get('max_rows', 100),
            False
        )
        result['_source'] = str(file_path)
        return result


# Tool registry
MSDOC_TOOLS = {
    'msdoc_list_files': MsDocListFilesTool,
    'msdoc_read_word': MsDocReadWordTool,
    'msdoc_read_excel': MsDocReadExcelTool,
    'msdoc_search_word': MsDocSearchWordTool,
    'msdoc_search_excel': MsDocSearchExcelTool,
    'msdoc_get_word_metadata': MsDocGetWordMetadataTool,
    'msdoc_get_excel_metadata': MsDocGetExcelMetadataTool,
    'msdoc_extract_text': MsDocExtractTextTool,
    'msdoc_get_excel_sheets': MsDocGetExcelSheetsTool,
    'msdoc_read_excel_sheet': MsDocReadExcelSheetTool
}
