"""
Shared base helpers for Microsoft Office document tools — upstream-compliant.

Ported from sajhamcpserver/sajha/tools/impl/msdoc_tools_tool_refactored.py.

Compliance changes from the original:
- No reliance on Flask `g`; worker context comes from `arguments['_worker_context']`.
- No `sajha.storage` — uses pathlib/open() for direct filesystem access.
- No `PropertiesConfigurator` fallback.
"""
import io
import os
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from sajha.tools.base_mcp_tool import BaseMCPTool

from tools_pack_lib.worker_ctx import get_data_layers


logger = logging.getLogger(__name__)

# Section parameter definition reused across all tool schemas
SECTION_PROP = {
    "type": "string",
    "enum": ["domain_data", "my_data", "common"],
    "description": (
        "Data layer: 'domain_data' (worker knowledge base, default), "
        "'my_data' (user uploads), 'common' (shared library)."
    ),
}


def _layer_roots(ctx: Dict[str, str], section: str = 'domain_data') -> List[str]:
    """Return list of candidate root paths for the requested section, primary first."""
    primary = []
    others = []
    for name, path in get_data_layers(ctx, 'all'):
        if name == 'domain_data' and section == 'domain_data':
            primary.append(path.rstrip('/'))
        elif name == 'my_data' and section == 'my_data':
            primary.append(path.rstrip('/'))
        elif name == 'common' and section == 'common':
            primary.append(path.rstrip('/'))
        else:
            others.append(path.rstrip('/'))
    return primary + others


class MsDocBaseTool(BaseMCPTool):
    """Base class for Microsoft Document tools with shared functionality."""

    def _resolve_docs_dir(self, ctx: Dict[str, str], section: str = 'domain_data') -> str:
        """Return the primary docs directory for the given section."""
        for name, path in get_data_layers(ctx, 'all'):
            if name == section and path:
                return path.rstrip('/')
        return ''

    def _candidate_paths(self, ctx: Dict[str, str], filename: str) -> List[str]:
        """Return all candidate absolute paths to search for a file, in priority order.

        Searches across all data layers so the tool finds the file regardless of
        which folder it was uploaded to:
          1. domain_data root + /msdocs/ subfolder (legacy)
          2. my_data
          3. common
        """
        candidates: List[str] = []
        for name, path in get_data_layers(ctx, 'all'):
            if not path:
                continue
            base = path.rstrip('/')
            candidates.append(os.path.join(base, filename))
            if name == 'domain_data':
                candidates.append(os.path.join(base, 'msdocs', filename))
        return candidates

    def _get_file_path(self, ctx: Dict[str, str], filename: str, section: str = 'domain_data') -> str:
        """Get full absolute path for a file, searching across all data folders."""
        # Try the requested section's primary dir first
        primary_root = self._resolve_docs_dir(ctx, section)
        if primary_root:
            primary = os.path.join(primary_root, filename)
            if os.path.exists(primary):
                return primary

        # Fall back to all candidate paths
        for path in self._candidate_paths(ctx, filename):
            if os.path.exists(path):
                return path

        # Last resort: recursive walk for bare filenames across all roots
        bare = os.path.basename(filename)
        for _, root in get_data_layers(ctx, 'all'):
            if not root:
                continue
            try:
                for dirpath, _, files in os.walk(root):
                    if bare in files:
                        return os.path.join(dirpath, bare)
            except OSError:
                continue

        # Nothing found — return primary path so callers get a meaningful "not found" error
        return os.path.join(primary_root or '.', filename)

    def _read_bytes(self, file_path: str) -> bytes:
        """Read a file as bytes — pure pathlib, no storage abstraction."""
        with open(file_path, 'rb') as fh:
            return fh.read()

    def _exists(self, file_path: str) -> bool:
        return bool(file_path) and os.path.exists(file_path)

    def _extract_word_heading(self, doc, heading_query: str):
        """Extract the content block under a matching Word heading style."""
        query = heading_query.strip().lower()
        from docx.table import Table as _DocxTable
        from docx.text.paragraph import Paragraph as _DocxPara

        body_children = list(doc.element.body)

        # Pass 1: index all heading paragraphs
        heading_map = []  # list of (child_index, level, text)
        for ci, child in enumerate(body_children):
            if child.tag.endswith('}p'):
                para = _DocxPara(child, doc)
                style_name = para.style.name if para.style else ''
                if style_name.startswith('Heading'):
                    text = para.text.strip()
                    if not text:
                        continue
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    heading_map.append((ci, level, text))

        if not heading_map:
            return None, None, []

        # Pass 2: find match
        match_i = None
        match_level = None
        match_title = None
        for hi, (ci, level, text) in enumerate(heading_map):
            if query in text.lower():
                match_i = hi
                match_level = level
                match_title = text
                break

        if match_i is None:
            return None, None, [h[2] for h in heading_map]

        # Pass 3: find end boundary
        start_ci = heading_map[match_i][0]
        end_ci = len(body_children)
        for hi in range(match_i + 1, len(heading_map)):
            next_ci, next_level, _ = heading_map[hi]
            if next_level <= match_level:
                end_ci = next_ci
                break

        # Pass 4: collect text from start_ci to end_ci
        lines = []
        for child in body_children[start_ci:end_ci]:
            if child.tag.endswith('}p'):
                para = _DocxPara(child, doc)
                if para.text.strip():
                    lines.append(para.text)
            elif child.tag.endswith('}tbl'):
                table = _DocxTable(child, doc)
                for row in table.rows:
                    row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_cells:
                        lines.append(' | '.join(row_cells))

        return '\n'.join(lines), match_title, [h[2] for h in heading_map]

    def _list_files_by_type(self, file_type: str, docs_dir: str) -> List[Dict]:
        """List files by type, walking subfolders."""
        files: List[Dict] = []
        if not docs_dir or not os.path.isdir(docs_dir):
            return files
        try:
            for dirpath, _, file_names in os.walk(docs_dir):
                for f in file_names:
                    if f.startswith('.'):
                        continue
                    extension = os.path.splitext(f)[1].lower()
                    if not extension:
                        continue
                    if file_type == 'word' and extension not in ['.docx', '.doc']:
                        continue
                    elif file_type == 'excel' and extension not in ['.xlsx', '.xls', '.xlsm']:
                        continue
                    elif file_type == 'all' and extension not in ['.docx', '.doc', '.xlsx', '.xls', '.xlsm']:
                        continue
                    abs_path = os.path.join(dirpath, f)
                    rel_path = os.path.relpath(abs_path, docs_dir)
                    try:
                        stat = os.stat(abs_path)
                        size = stat.st_size
                        modified = stat.st_mtime
                    except OSError:
                        size = 0
                        modified = 0.0
                    files.append({
                        'filename': rel_path,
                        'path': abs_path,
                        'extension': extension,
                        'size': size,
                        'modified': modified,
                    })
            return sorted(files, key=lambda x: x['modified'], reverse=True)
        except Exception as e:
            logger.error(f"Failed to list files: {e}")
            return []

    def _read_word_document(self, file_path: str) -> Dict:
        """Read Word document via direct filesystem access."""
        try:
            from docx import Document
            raw = self._read_bytes(file_path)
            doc = Document(io.BytesIO(raw))

            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            return {
                'filename': os.path.basename(file_path),
                'paragraphs': paragraphs,
                'paragraph_count': len(paragraphs),
                'tables': tables,
                'table_count': len(tables)
            }

        except ImportError:
            raise ValueError("python-docx library not installed. Install with: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Failed to read Word document: {str(e)}")

    def _read_excel_document(self, file_path: str, sheet_name: Optional[str] = None,
                             sheet_index: Optional[int] = None, max_rows: int = 100,
                             include_formulas: bool = False) -> Dict:
        """Read Excel document via direct filesystem access."""
        try:
            from openpyxl import load_workbook

            raw = self._read_bytes(file_path)
            wb = load_workbook(io.BytesIO(raw), data_only=not include_formulas)

            if sheet_name:
                sheet = wb[sheet_name]
            elif sheet_index is not None:
                sheet = wb.worksheets[sheet_index]
            else:
                sheet = wb.active

            data = []
            for i, row in enumerate(sheet.iter_rows(values_only=not include_formulas)):
                if i >= max_rows:
                    break
                data.append(list(row))

            result = {
                'filename': os.path.basename(file_path),
                'sheet_name': sheet.title,
                'data': data,
                'row_count': len(data),
                'column_count': len(data[0]) if data else 0
            }

            if include_formulas:
                formulas = []
                for row in sheet.iter_rows():
                    row_formulas = []
                    for cell in row:
                        if cell.value and isinstance(cell.value, str) and cell.value.startswith('='):
                            row_formulas.append({
                                'cell': cell.coordinate,
                                'formula': cell.value
                            })
                    if row_formulas:
                        formulas.append(row_formulas)
                result['formulas'] = formulas

            wb.close()
            return result

        except ImportError:
            raise ValueError("openpyxl library not installed. Install with: pip install openpyxl")
        except Exception as e:
            raise ValueError(f"Failed to read Excel document: {str(e)}")
