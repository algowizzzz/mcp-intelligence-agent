"""
md_to_docx — compliant with upstream SAJHA.

Ported from sajhamcpserver/sajha/tools/impl/operational_tools.py (MdToDocxTool).
Worker scope from arguments['_worker_context'].
"""

import logging
import os
import re
from pathlib import Path
from typing import Any, Dict

from sajha.tools.base_mcp_tool import BaseMCPTool

from tools_pack_impl._operational_base import get_roots, safe_path, find_file

logger = logging.getLogger(__name__)


def _add_inline(para, text):
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            from docx.shared import Pt
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            para.add_run(part)


class MdToDocx(BaseMCPTool):
    """Convert a .md file to a formatted .docx Word document."""

    def get_input_schema(self) -> Dict:
        return self.config.get('inputSchema', {
            'type': 'object',
            'properties': {
                'file_path':   {'type': 'string'},
                'output_path': {'type': 'string'},
                'style':       {'type': 'string', 'enum': ['standard', 'minimal', 'report']},
                'include_toc': {'type': 'boolean'},
            },
            'required': ['file_path'],
        })

    def get_output_schema(self) -> Dict:
        return {'type': 'object'}

    def execute(self, arguments: Dict[str, Any]) -> Dict[str, Any]:
        file_path = arguments.get('file_path', '')
        output_path = arguments.get('output_path')
        include_toc = arguments.get('include_toc', False)

        domain, my_data, common = get_roots(arguments)
        safe = safe_path(file_path, domain, my_data, common)
        if not safe or not safe.exists():
            safe = find_file(file_path, domain, my_data, common)
        if not safe or not safe.exists():
            return {'error': f'File not found or access denied: {file_path}'}
        if safe.suffix.lower() != '.md':
            return {'error': 'md_to_docx only accepts .md files.'}

        text = safe.read_text(encoding='utf-8', errors='replace')
        text = re.sub(r'^---\n.*?\n---\n', '', text, flags=re.DOTALL)

        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError:
            return {'error': 'python-docx is not installed.'}

        doc = Document()
        for section in doc.sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        lines = text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            m = re.match(r'^(#{1,6})\s+(.*)', line)
            if m:
                level = min(len(m.group(1)), 3)
                doc.add_heading(m.group(2).strip(), level=level)
                i += 1
                continue
            if re.match(r'^---+$', line.strip()):
                doc.add_paragraph()
                i += 1
                continue
            if line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                i += 1
                para = doc.add_paragraph('\n'.join(code_lines))
                para.style = 'Normal'
                for run in para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                continue
            if '|' in line and line.strip().startswith('|'):
                table_rows = []
                while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                    row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                    if not all(re.match(r'^[-:]+$', c) for c in row if c):
                        table_rows.append(row)
                    i += 1
                if table_rows:
                    max_cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=max_cols)
                    table.style = 'Table Grid'
                    for ri, row in enumerate(table_rows):
                        for ci, cell_text in enumerate(row):
                            if ci < max_cols:
                                cell = table.rows[ri].cells[ci]
                                cell.text = cell_text
                                if ri == 0:
                                    for run in cell.paragraphs[0].runs:
                                        run.font.bold = True
                continue
            if re.match(r'^[\-\*\+]\s+', line):
                text_part = re.sub(r'^[\-\*\+]\s+', '', line)
                para = doc.add_paragraph(style='List Bullet')
                _add_inline(para, text_part)
                i += 1
                continue
            if re.match(r'^\d+\.\s+', line):
                text_part = re.sub(r'^\d+\.\s+', '', line)
                para = doc.add_paragraph(style='List Number')
                _add_inline(para, text_part)
                i += 1
                continue
            if line.startswith('> '):
                para = doc.add_paragraph(line[2:])
                para.paragraph_format.left_indent = Inches(0.5)
                i += 1
                continue
            if not line.strip():
                i += 1
                continue
            para = doc.add_paragraph()
            _add_inline(para, line)
            i += 1

        if output_path:
            out = Path(output_path)
        else:
            out = safe.with_suffix('.docx')
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        size_bytes = out.stat().st_size if out.exists() else 0

        return {
            'source_md':    str(safe),
            'output_docx':  str(out),
            'size_bytes':   size_bytes,
            'style':        arguments.get('style', 'standard'),
            'toc_included': include_toc,
            '_source':      str(out),
        }
