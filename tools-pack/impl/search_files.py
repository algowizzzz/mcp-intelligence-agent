"""
search_files — compliant with upstream SAJHA.

Ported from sajhamcpserver/sajha/tools/impl/operational_tools.py
(SearchFilesTool). Worker scope from arguments['_worker_context'].
"""

import io
import logging
import os
import re
from typing import Any, Dict, List

from sajha.tools.base_mcp_tool import BaseMCPTool

from tools_pack_impl._operational_base import get_roots

logger = logging.getLogger(__name__)


def _extract_text(path: str, ext: str) -> str:
    try:
        if ext in ('md', 'txt', 'csv', 'tsv', 'json'):
            with open(path, 'r', encoding='utf-8', errors='replace') as fh:
                return fh.read()
        if ext == 'docx':
            from docx import Document
            with open(path, 'rb') as fh:
                doc = Document(io.BytesIO(fh.read()))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return '\n'.join(parts)
        if ext in ('xlsx', 'xls'):
            import openpyxl
            with open(path, 'rb') as fh:
                wb = openpyxl.load_workbook(io.BytesIO(fh.read()), read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    parts.append(' '.join(str(c) for c in row if c is not None))
            return '\n'.join(parts)
        if ext == 'pdf':
            try:
                import fitz
                with open(path, 'rb') as fh:
                    doc = fitz.open(stream=io.BytesIO(fh.read()), filetype='pdf')
                return '\n'.join(doc[i].get_text() for i in range(len(doc)))
            except Exception:
                return ''
    except Exception:
        return ''
    return ''


class SearchFiles(BaseMCPTool):
    """Full-text keyword search across worker data layers."""

    def get_input_schema(self) -> Dict:
        return self.config.get('inputSchema', {
            'type': 'object',
            'properties': {
                'query':         {'type': 'string'},
                'section':       {'type': 'string', 'enum': ['domain_data', 'my_data', 'common', 'all']},
                'file_type':     {'type': 'string'},
                'folder':        {'type': 'string'},
                'max_results':   {'type': 'integer'},
                'excerpt_chars': {'type': 'integer'},
            },
            'required': ['query'],
        })

    def get_output_schema(self) -> Dict:
        return {'type': 'object'}

    def execute(self, arguments: Dict[str, Any]) -> Dict[str, Any]:
        query = arguments.get('query', '')
        section = arguments.get('section', 'all')
        file_type = arguments.get('file_type', 'all')
        folder_filter = (arguments.get('folder') or '').strip()
        max_results = int(arguments.get('max_results', 20))
        excerpt_chars = int(arguments.get('excerpt_chars', 200))

        if not query:
            return {'error': 'query is required', 'results': []}

        exact = re.match(r'^"(.+)"$', query)
        if exact:
            pattern = re.compile(re.escape(exact.group(1)), re.IGNORECASE)
        else:
            pattern = re.compile(re.escape(query), re.IGNORECASE)

        domain, my_data, common = get_roots(arguments)
        layer_map = {'domain_data': domain, 'my_data': my_data, 'common': common}

        roots: List = []
        if section == 'all':
            for n, r in layer_map.items():
                if r:
                    roots.append((n, str(r)))
        elif section in layer_map and layer_map[section]:
            roots.append((section, str(layer_map[section])))

        candidates: List = []
        for sec_name, root_str in roots:
            if not os.path.isdir(root_str):
                continue
            for dirpath, _, filenames in os.walk(root_str):
                for fname in filenames:
                    if fname.startswith('.'):
                        continue
                    rel_path = os.path.relpath(os.path.join(dirpath, fname), root_str)
                    if folder_filter and folder_filter.lower() not in rel_path.lower():
                        continue
                    ext = os.path.splitext(fname)[1].lower().lstrip('.')
                    if file_type != 'all' and ext != file_type:
                        continue
                    if ext in ('parquet', 'pq', 'db', 'wal', 'pyc'):
                        continue
                    f_path = os.path.join(root_str, rel_path)
                    candidates.append((sec_name, f_path, fname, ext))

        results: List[Dict[str, Any]] = []
        for sec_name, f_path, fname, ext in candidates:
            if len(results) >= max_results:
                break
            text = _extract_text(f_path, ext)
            if not text:
                continue
            matches = list(pattern.finditer(text))
            if not matches:
                continue
            excerpts = []
            for m in matches[:5]:
                start = max(0, m.start() - excerpt_chars // 2)
                end = min(len(text), m.end() + excerpt_chars // 2)
                snippet = text[start:end].replace('\n', ' ').strip()
                snippet = pattern.sub(lambda x: f"[{x.group()}]", snippet)
                excerpts.append(snippet)
            results.append({
                'filename':    fname,
                'path':        f_path,
                'file_type':   ext,
                'section':     sec_name,
                'match_count': len(matches),
                'excerpts':    excerpts,
                '_source':     f_path,
            })

        return {
            'query':         query,
            'total_matches': len(results),
            'results':       results,
        }
