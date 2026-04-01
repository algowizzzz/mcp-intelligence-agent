"""
Script to write two .docx files:
1. Sajha_IR_Tavily_Replacement_ERD.docx  (updated/overwrite)
2. Tavily_SEC_MDA_Direct_Fetch_Note.docx  (new)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_DIR = "/Users/saadahmed/Desktop/react_agent/requirements"

# ─────────────────────────────────────────────────────────────────────────────
# HELPER UTILITIES
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table_row(table, values, bold=False, bg_hex=None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = bold
                run.font.size = Pt(9)
        if bg_hex:
            set_cell_bg(cell, bg_hex)
    return row


def add_code_para(doc, text):
    """Add a paragraph styled as monospace code."""
    p = doc.add_paragraph(text)
    p.style = doc.styles["No Spacing"]
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def make_header_row(table, *headers):
    """Populate the first row of a table as a bold header row."""
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        for para in hdr.cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_bg(hdr.cells[i], "D9E1F2")


# ─────────────────────────────────────────────────────────────────────────────
# FILE 1 — Sajha_IR_Tavily_Replacement_ERD.docx
# ─────────────────────────────────────────────────────────────────────────────

def build_erd():
    doc = Document()

    # ── narrow margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Title block ─────────────────────────────────────────────────────────
    t1 = doc.add_paragraph("ENGINEERING REQUIREMENTS")
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1.runs[0].bold = True
    t1.runs[0].font.size = Pt(18)

    t2 = doc.add_paragraph("Sajha IR Tools — Tavily Replacement")
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].bold = True
    t2.runs[0].font.size = Pt(14)
    t2.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    t3 = doc.add_paragraph(
        "Deprecate web-scraper stack \u00b7 Replace with 3 Tavily-native tools \u00b7 Universal company coverage"
    )
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.runs[0].font.size = Pt(10)
    t3.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    t4 = doc.add_paragraph("Version 1.1  \u00b7  March 2026")
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t4.runs[0].font.size = Pt(9)
    t4.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # 1. Overview
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Overview", level=1)

    doc.add_paragraph(
        "The existing Investor Relations (IR) tool stack on Sajha relies on a custom web-scraper "
        "framework (sajha/ir/) that maintains a hardcoded database of 10 companies, navigates their "
        "IR pages via HTML parsing, and returns document URLs with no content extraction. This "
        "architecture is brittle, limited in coverage, and cannot read document content."
    )
    doc.add_paragraph(
        "This document specifies the full replacement of that stack with three Tavily-native tools "
        "that work for any publicly traded company worldwide, discover documents dynamically through "
        "search, and extract actual document content using Tavily\u2019s extract endpoint \u2014 "
        "capability the current scraper-based tools do not have at all."
    )

    # 1.1 Objectives
    doc.add_heading("1.1 Objectives", level=2)
    objectives = [
        "Remove all 7 existing IR tool JSON configs and the investor_relations_tool_refactored.py implementation file.",
        "Remove the entire sajha/ir/ module (20 files: scrapers, factory, company database, HTTP client, SEC Edgar client, documentation).",
        "Replace with 3 Tavily-native tools: ir_find_page, ir_find_documents, ir_extract_content.",
        "Achieve universal company coverage \u2014 any publicly traded company, no hardcoded database.",
        "Enable actual document content extraction, not just link discovery.",
        "Apply a simple size guardrail: reject responses exceeding 100KB with a clear error message.",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Bullet")

    # 1.2 Scope
    doc.add_heading("1.2 Scope", level=2)
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Table Grid"
    make_header_row(tbl, "In Scope", "Out of Scope")
    scope_data = [
        ("7 IR tool JSON configs (config/tools/)",
         "IRIS CCR tools (iris_ccr_tools.py) \u2014 these are separate internal risk tools and are NOT affected"),
        ("investor_relations_tool_refactored.py",
         "EDGAR tools \u2014 covered under separate EDGAR architecture plan"),
        ("Entire sajha/ir/ directory (20 files)",
         "Yahoo Finance tools \u2014 no changes"),
        ("3 new Tavily IR tool JSON configs",
         "Tavily domain search config \u2014 no changes needed"),
        ("1 new Tavily IR tool impl file", ""),
    ]
    for in_s, out_s in scope_data:
        row = tbl.add_row()
        row.cells[0].text = in_s
        row.cells[1].text = out_s
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # ════════════════════════════════════════════════════════════════════════
    # 2. Current State
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Current State \u2014 Tools Being Removed", level=1)

    # 2.1 Problems
    doc.add_heading("2.1 Problems with the Existing Architecture", level=2)
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = "Table Grid"
    make_header_row(tbl2, "Problem", "Impact")
    problems = [
        ("Hardcoded company database \u2014 only 10 companies supported (AAPL, MSFT, GOOGL, AMZN, NVDA, TSLA, JPM, GS, BAC, WMT)",
         "Any other ticker throws an error. Zero extensibility without code changes."),
        ("HTML scraper approach \u2014 parses anchor tags from IR pages",
         "Breaks whenever a company redesigns its IR site. No maintenance mechanism exists."),
        ("Returns URLs only \u2014 no document content extraction",
         "The LLM receives a list of links and cannot answer questions about what the documents contain."),
        ("SEC EDGAR fallback calls companyfacts endpoint (5MB dump)",
         "Same context overflow issue identified in the EDGAR architecture review."),
        ("Auto-discovery not implemented \u2014 returns None for unknown tickers",
         "The auto-discovery code path exists but does nothing. 10 companies is the hard ceiling."),
        ("20-file module with no tests",
         "High maintenance burden. The sajha/ir/ module is larger than the tools it supports."),
    ]
    for prob, imp in problems:
        row = tbl2.add_row()
        row.cells[0].text = prob
        row.cells[1].text = imp
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # 2.2 Files to Delete
    doc.add_heading("2.2 Files to Delete", level=2)

    doc.add_heading("2.2.1 JSON Tool Configs (config/tools/)", level=3)
    doc.add_paragraph("Delete the following 7 files from sajhamcpserver/config/tools/:")
    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = "Table Grid"
    make_header_row(tbl3, "File", "Tool Name", "Replacement")
    json_files = [
        ("ir_find_page.json", "ir_find_page", "Replaced by new ir_find_page.json (Tavily-based)"),
        ("ir_get_all_resources.json", "ir_get_all_resources", "Replaced by ir_find_documents.json"),
        ("ir_get_annual_reports.json", "ir_get_annual_reports", "Replaced by ir_find_documents.json"),
        ("ir_get_documents.json", "ir_get_documents", "Replaced by ir_find_documents.json"),
        ("ir_get_latest_earnings.json", "ir_get_latest_earnings", "Replaced by ir_find_documents.json"),
        ("ir_get_presentations.json", "ir_get_presentations", "Replaced by ir_find_documents.json"),
        ("ir_list_supported_companies.json", "ir_list_supported_companies",
         "Not replaced \u2014 concept eliminated (universal coverage)"),
    ]
    for f, t, r in json_files:
        row = tbl3.add_row()
        row.cells[0].text = f
        row.cells[1].text = t
        row.cells[2].text = r
        for c in row.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_heading("2.2.2 Implementation File (sajha/tools/impl/)", level=3)
    doc.add_paragraph("Delete the following file from sajhamcpserver/sajha/tools/impl/:")
    doc.add_paragraph("investor_relations_tool_refactored.py   \u2014 all 7 tool classes, 555 lines")

    doc.add_heading("2.2.3 IR Scraper Module (sajha/ir/)", level=3)
    doc.add_paragraph("Delete the entire sajha/ir/ directory and all its contents (20 files):")
    tbl4 = doc.add_table(rows=1, cols=2)
    tbl4.style = "Table Grid"
    make_header_row(tbl4, "File", "Description")
    ir_files = [
        ("__init__.py", "Module init"),
        ("base_ir_webscraper.py", "Abstract base class for all scrapers"),
        ("enhanced_base_scraper.py", "Enhanced base with link extraction utilities"),
        ("generic_ir_scraper.py", "Generic HTML scraper for any IR page"),
        ("company_ir_scrapers.py", "Company-specific scraper overrides"),
        ("ir_webscraper_factory.py", "Original scraper factory"),
        ("enhanced_factory.py", "Enhanced factory with auto-discovery (stub)"),
        ("company_database.py", "Hardcoded 10-company configuration database"),
        ("http_client.py", "HTTP fetch utilities for scraping"),
        ("sec_edgar.py", "SEC EDGAR client (companyfacts fallback)"),
        ("demo.py", "Demo/test script"),
        ("README.md", "Module documentation"),
        ("FINAL_SUMMARY.md", "Architecture summary doc"),
        ("INTEGRATION_GUIDE.md", "Integration guide"),
        ("MIGRATION.md", "Migration notes"),
        ("QUICK_START.md", "Quick start guide"),
        ("Investor Relations Tool - Architecture Diagram.md", "Architecture diagram"),
        ("Investor Relations Tool - Complete Package Index.md", "Package index"),
        ("Investor Relations Tool - Executive Summary.md", "Executive summary"),
        ("Investor Relations Tool - Testing Guide.md", "Testing guide"),
    ]
    for fname, desc in ir_files:
        row = tbl4.add_row()
        row.cells[0].text = fname
        row.cells[1].text = desc
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph("Removal command:")
    for cmd in [
        "rm -rf sajhamcpserver/sajha/ir/",
        "rm sajhamcpserver/sajha/tools/impl/investor_relations_tool_refactored.py",
        "rm sajhamcpserver/config/tools/ir_find_page.json",
        "rm sajhamcpserver/config/tools/ir_get_all_resources.json",
        "rm sajhamcpserver/config/tools/ir_get_annual_reports.json",
        "rm sajhamcpserver/config/tools/ir_get_documents.json",
        "rm sajhamcpserver/config/tools/ir_get_latest_earnings.json",
        "rm sajhamcpserver/config/tools/ir_get_presentations.json",
        "rm sajhamcpserver/config/tools/ir_list_supported_companies.json",
    ]:
        add_code_para(doc, cmd)

    # ════════════════════════════════════════════════════════════════════════
    # 3. New Architecture
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("3. New Architecture \u2014 Tavily-First IR", level=1)

    doc.add_heading("3.1 Design Principles", level=2)
    principles = [
        "Search-first: Tavily search discovers document URLs dynamically for any company \u2014 no database, no scraper.",
        "Extract on demand: Tavily extract reads document content (HTML and PDF) when called, returning relevant text via the query parameter.",
        "Query-focused: The user query is always passed to Tavily extract so it returns only the relevant portion of large documents \u2014 no custom chunking or secondary LLM calls needed.",
        "Universal coverage: Works for any public company worldwide, not just hardcoded tickers.",
        "Size guardrail: A single byte-length check rejects responses over 100KB before they reach the LLM, with a descriptive error message.",
        "Same pattern as Yahoo Finance: Follows the proven Tavily+LLM architecture already in tavily_yahoo_finance_tool.py.",
    ]
    for p in principles:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("3.2 Agent Flow", level=2)
    doc.add_paragraph(
        "For a query such as \u201cWhat was Apple\u2019s revenue guidance in the Q4 2024 earnings call?\u201d, "
        "the agent flow is:"
    )
    tbl_flow = doc.add_table(rows=1, cols=4)
    tbl_flow.style = "Table Grid"
    make_header_row(tbl_flow, "Step", "Tool Called", "Tavily Operation", "Output")
    flow_rows = [
        ("1", "ir_find_page", 'search: "Apple Inc investor relations official site"',
         "ir_url: https://investor.apple.com"),
        ("2", "ir_find_documents",
         'search: "Apple Q4 2024 earnings press release site:investor.apple.com OR site:sec.gov"',
         "List of {title, url, date, type}"),
        ("3", "ir_extract_content", 'extract: url + query="Q4 2024 revenue guidance"',
         "Relevant text passages from the document"),
        ("4", "LLM", "\u2014", "Synthesised answer from the extracted content"),
    ]
    for vals in flow_rows:
        row = tbl_flow.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_heading("3.3 Tool Summary", level=2)
    tbl_sum = doc.add_table(rows=1, cols=3)
    tbl_sum.style = "Table Grid"
    make_header_row(tbl_sum, "Tool", "Purpose", "Replaces")
    tool_summary = [
        ("ir_find_page",
         "Find the canonical investor relations page URL for any company",
         "ir_find_page.json + company_database.py + EnhancedIRScraperFactory"),
        ("ir_find_documents",
         "Search for specific document types (annual report, earnings, presentation, proxy, press release) for a company",
         "ir_get_documents, ir_get_annual_reports, ir_get_presentations, ir_get_latest_earnings, ir_get_all_resources"),
        ("ir_extract_content",
         "Extract relevant text from a document URL using Tavily extract with query-focused filtering",
         "Nothing \u2014 this capability did not exist in the old stack"),
    ]
    for vals in tool_summary:
        row = tbl_sum.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # ════════════════════════════════════════════════════════════════════════
    # 4. Tool Specifications
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Tool Specifications", level=1)

    # ── 4.1 ir_find_page ────────────────────────────────────────────────────
    doc.add_heading("4.1 ir_find_page", level=2)
    doc.add_paragraph(
        "Discovers the official investor relations page URL for any publicly traded company using "
        "Tavily search. This provides the canonical IR domain used to scope subsequent document searches."
    )

    doc.add_heading("Input Schema", level=3)
    tbl_in1 = doc.add_table(rows=1, cols=4)
    tbl_in1.style = "Table Grid"
    make_header_row(tbl_in1, "Parameter", "Type", "Required", "Description")
    row = tbl_in1.add_row()
    row.cells[0].text = "company"
    row.cells[1].text = "string"
    row.cells[2].text = "Yes"
    row.cells[3].text = 'Company name or ticker symbol (e.g. "Apple", "AAPL", "Tesla Inc")'
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

    doc.add_heading("Output Schema", level=3)
    tbl_out1 = doc.add_table(rows=1, cols=3)
    tbl_out1.style = "Table Grid"
    make_header_row(tbl_out1, "Field", "Type", "Description")
    out1_rows = [
        ("ticker", "string", "Resolved ticker symbol if identifiable"),
        ("company_name", "string", "Full company name as resolved by Tavily"),
        ("ir_url", "string", "Canonical investor relations page URL"),
        ("source", "string", "Source URL of the result"),
    ]
    for vals in out1_rows:
        row = tbl_out1.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_heading("Tavily Call Pattern", level=3)
    for line in [
        "tavily.search(",
        '  query = f"{company} investor relations official site",',
        '  search_depth = "basic",',
        "  max_results = 3",
        ")",
    ]:
        add_code_para(doc, line)

    # CHANGE 5: expanded aggregator filter list
    doc.add_paragraph(
        "Skip results where the URL domain matches any known financial aggregator: "
        "sec.gov, yahoo.com, finance.yahoo.com, morningstar.com, macrotrends.net, "
        "stockanalysis.com, wsj.com, bloomberg.com, reuters.com, wisesheets.io, "
        "simplywall.st, finance.google.com, marketwatch.com, investing.com, "
        "seekingalpha.com, zacks.com, fool.com. "
        "Also add a positive signal check: prefer results where the domain matches a known IR "
        "hosting platform (q4inc.com, q4web.com, workiva.com, businesswire.com, prnewswire.com) "
        "OR contains the company name."
    )

    doc.add_heading("JSON Config Location", level=3)
    doc.add_paragraph("sajhamcpserver/config/tools/ir_find_page.json")

    doc.add_heading("Implementation Class", level=3)
    doc.add_paragraph("File:  sajhamcpserver/sajha/tools/impl/tavily_ir_tool.py")
    doc.add_paragraph("Class: IRFindPageTool")

    # ── 4.2 ir_find_documents ───────────────────────────────────────────────
    doc.add_heading("4.2 ir_find_documents", level=2)
    doc.add_paragraph(
        "Searches for investor relations documents of a specified type for a given company. Uses "
        "Tavily search scoped to the company\u2019s IR domain and SEC EDGAR to return ranked results "
        "with titles, URLs, and dates. Replaces five separate scraper-based tools with a single "
        "parameterised search tool."
    )

    doc.add_heading("Input Schema", level=3)
    tbl_in2 = doc.add_table(rows=1, cols=5)
    tbl_in2.style = "Table Grid"
    make_header_row(tbl_in2, "Parameter", "Type", "Required", "Description", "Example Values")
    in2_rows = [
        ("company", "string", "Yes", "Company name or ticker", '"Apple", "AAPL"'),
        ("document_type", "string", "Yes", "Category of document to find",
         '"annual_report", "quarterly_earnings", "earnings_presentation", "investor_presentation", '
         '"proxy_statement", "press_release", "esg_report", "sec_filing", "all"'),
        ("year", "integer", "No", "Filter to a specific year", "2024, 2023"),
        ("quarter", "string", "No", "Filter to a specific quarter", '"Q1", "Q2", "Q3", "Q4"'),
        ("ir_url", "string", "No", "IR domain from ir_find_page (improves precision)",
         '"https://investor.apple.com"'),
        ("max_results", "integer", "No", "Max documents to return. Default 5, max 10", "5"),
    ]
    for vals in in2_rows:
        row = tbl_in2.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_heading("Output Schema", level=3)
    tbl_out2 = doc.add_table(rows=1, cols=3)
    tbl_out2.style = "Table Grid"
    make_header_row(tbl_out2, "Field", "Type", "Description")
    out2_rows = [
        ("company", "string", "Company name or ticker as provided"),
        ("document_type", "string", "Document type searched"),
        ("count", "integer", "Number of results returned"),
        ("documents", "array", "Array of {title, url, date, type, snippet}"),
        ("documents[].title", "string", "Document title as indexed by Tavily"),
        ("documents[].url", "string", "Direct URL to the document"),
        ("documents[].date", "string", "Publication date if available"),
        ("documents[].snippet", "string", "Short excerpt from the document page"),
    ]
    for vals in out2_rows:
        row = tbl_out2.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_heading("Document Type to Query Mapping", level=3)
    tbl_map = doc.add_table(rows=1, cols=2)
    tbl_map.style = "Table Grid"
    make_header_row(tbl_map, "document_type Value", "Tavily Search Query Pattern")
    map_rows = [
        ("annual_report", '"{company} annual report 10-K {year} site:{ir_domain} OR site:sec.gov"'),
        ("quarterly_earnings",
         '"{company} Q{quarter} {year} earnings press release site:{ir_domain} OR site:sec.gov"'),
        ("earnings_presentation",
         '"{company} Q{quarter} {year} earnings presentation slides site:{ir_domain}"'),
        ("investor_presentation",
         '"{company} investor day presentation {year} site:{ir_domain}"'),
        ("proxy_statement", '"{company} proxy statement DEF 14A {year} site:sec.gov"'),
        ("press_release", '"{company} press release {year} site:{ir_domain}"'),
        ("esg_report", '"{company} ESG sustainability report {year} site:{ir_domain}"'),
        ("sec_filing", '"{company} SEC filing {year} site:sec.gov"'),
        ("all",
         '"{company} investor relations documents {year} site:{ir_domain} OR site:sec.gov"'),
    ]
    for vals in map_rows:
        row = tbl_map.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # CHANGE 4: sec_filing note
    p_sec = doc.add_paragraph()
    run_bold = p_sec.add_run("sec_filing type: ")
    run_bold.bold = True
    run_bold.font.size = Pt(9)
    run_rest = p_sec.add_run(
        "Using tavily_search on sec.gov for structured SEC filings is unreliable (returns wrong "
        "companies, stale documents). When document_type='sec_filing', the tool should delegate to "
        "the EDGAR _resolve_filing_url() pipeline (direct SEC submissions API) rather than keyword "
        "search. This is the same fix applied to all EDGAR qualitative tools. "
        "Reference: edgar_tavily_tools.py _resolve_filing_url()."
    )
    run_rest.font.size = Pt(9)

    # CHANGE 7: fallback query note
    p_fallback = doc.add_paragraph()
    run_bold2 = p_fallback.add_run("Fallback when ir_url is absent: ")
    run_bold2.bold = True
    run_bold2.font.size = Pt(9)
    run_rest2 = p_fallback.add_run(
        "When ir_url is not provided (caller skipped ir_find_page), fall back to an open search "
        "without site scoping: '{company} {document_type_keywords} {year} investor relations' with "
        "search_depth='advanced'. Do not fail with an error when ir_url is absent."
    )
    run_rest2.font.size = Pt(9)

    doc.add_heading("JSON Config Location", level=3)
    doc.add_paragraph("sajhamcpserver/config/tools/ir_find_documents.json")

    doc.add_heading("Implementation Class", level=3)
    doc.add_paragraph("File:  sajhamcpserver/sajha/tools/impl/tavily_ir_tool.py")
    doc.add_paragraph("Class: IRFindDocumentsTool")

    # ── 4.3 ir_extract_content ──────────────────────────────────────────────
    doc.add_heading("4.3 ir_extract_content", level=2)
    doc.add_paragraph(
        "Extracts and structures content from an investor relations document URL. Uses Tavily extract "
        "to fetch the document, then passes the raw content to llm_extract() which makes a secondary "
        "Anthropic API call with a targeted extraction prompt. This two-step pattern is identical to "
        "all EDGAR tools and ensures the main LLM receives compact structured JSON, not raw HTML or "
        "PDF text."
    )
    doc.add_paragraph(
        "This is the capability that did not exist in the previous IR tool stack. The old tools "
        "returned document URLs and stopped. This tool reads what is inside."
    )

    doc.add_heading("Input Schema", level=3)
    tbl_in3 = doc.add_table(rows=1, cols=4)
    tbl_in3.style = "Table Grid"
    make_header_row(tbl_in3, "Parameter", "Type", "Required", "Description")
    in3_rows = [
        ("url", "string", "Yes",
         "Document URL to extract from (HTML page or PDF). Typically a URL returned by ir_find_documents."),
        ("query", "string", "Yes",
         "The user\u2019s question or topic. Passed to Tavily extract to filter returned content to relevant sections only."),
    ]
    for vals in in3_rows:
        row = tbl_in3.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_heading("Output Schema", level=3)
    tbl_out3 = doc.add_table(rows=1, cols=3)
    tbl_out3.style = "Table Grid"
    make_header_row(tbl_out3, "Field", "Type", "Description")
    out3_rows = [
        ("url", "string", "Source URL of the extracted document"),
        ("query", "string", "Query used for extraction"),
        ("content", "string", "Extracted relevant text from the document"),
        ("content_length_bytes", "integer", "Byte length of extracted content"),
        ("truncated", "boolean", "True if content was truncated by the size guardrail"),
    ]
    for vals in out3_rows:
        row = tbl_out3.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_heading("Execution Pattern", level=3)
    doc.add_paragraph("Follows the same two-step pattern as EDGAR tools (edgar_tavily_client.py):")
    code_lines = [
        "# Step 1: Fetch document content via Tavily extract",
        "raw_results = tavily_extract([url], query=query)   # pass query for focused extraction",
        'content = raw_results[0].get("raw_content", "")',
        "",
        "# Step 1a: Unescape underscores on SEC-hosted documents",
        'if "sec.gov" in url:',
        "    content = fix_tavily_json(content)       # Tavily escapes underscores in raw_content",
        "",
        "# Step 1b: Fallback if extract returns insufficient content",
        "if len(content.strip()) < 200:",
        "    # Return snippet already held by ir_find_documents for this URL",
        "    # (caller should pass snippet= from the ir_find_documents result)",
        '    content = snippet or ""',
        "",
        "# Step 2: Structured extraction via secondary LLM call",
        "result = llm_extract(content, extraction_prompt)",
        "# extraction_prompt is built from the user query — returns structured JSON",
    ]
    for line in code_lines:
        add_code_para(doc, line)

    # CHANGE 2: SEC Archives URL Handling
    p_sec_arch = doc.add_paragraph()
    run_bold_arch = p_sec_arch.add_run("SEC Archives URL Handling: ")
    run_bold_arch.bold = True
    run_bold_arch.font.size = Pt(10)
    run_arch = p_sec_arch.add_run(
        "When the document URL is from sec.gov/Archives (e.g. a 10-K or proxy filing returned by "
        "ir_find_documents), tavily_extract will fail for documents larger than ~2 MB. These filings "
        "can reach 12 MB. In this case, ir_extract_content must call "
        "stream_sec_section(url, section_marker) from edgar_tavily_client.py instead. Detection: if "
        "'sec.gov/Archives' in url, use stream_sec_section with the first keyword from query as the "
        "section_marker. This function is already implemented and tested."
    )
    run_arch.font.size = Pt(10)

    # CHANGE 6: Company name validation
    p_valid = doc.add_paragraph()
    run_bold_v = p_valid.add_run("Company name validation: ")
    run_bold_v.bold = True
    run_bold_v.font.size = Pt(10)
    run_v = p_valid.add_run(
        "Parallel to _validate_sources() in EDGAR tools: after llm_extract() returns, check whether "
        "the extracted content or the raw Tavily content mentions the company name "
        "(case-insensitive). If neither mentions it, set data_quality='WARN' and add a warning "
        "field: 'Extracted content does not reference {company}. Review sources manually.'"
    )
    run_v.font.size = Pt(10)

    doc.add_heading("Input Schema \u2014 Additional Field", level=3)
    tbl_add = doc.add_table(rows=1, cols=4)
    tbl_add.style = "Table Grid"
    make_header_row(tbl_add, "Parameter", "Type", "Required", "Description")
    add_rows = [
        ("url", "string", "Yes",
         "Document URL to extract from (HTML page or PDF). Typically a URL returned by ir_find_documents."),
        ("query", "string", "Yes",
         "The user\u2019s question or topic. Used to build the llm_extract() prompt targeting relevant content."),
        ("snippet", "string", "No",
         "Optional snippet text from ir_find_documents result for this URL. Used as fallback if Tavily extract returns insufficient content (<200 chars)."),
    ]
    for vals in add_rows:
        row = tbl_add.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_heading("Supported Document Formats", level=3)
    tbl_fmt = doc.add_table(rows=1, cols=3)
    tbl_fmt.style = "Table Grid"
    make_header_row(tbl_fmt, "Format", "Support", "Notes")
    fmt_rows = [
        ("HTML (IR pages, SEC EDGAR HTML filings)", "Full",
         "Best results. Tavily extract parses HTML structure cleanly. fix_tavily_json() applied for sec.gov URLs."),
        ("Text-based PDF (most modern SEC filings, earnings releases)", "Good",
         "Tavily can extract text from accessible PDFs. llm_extract() then structures the output."),
        ("Scanned / image PDF (older annual reports)", "Fallback",
         "Tavily returns empty or garbled text. Tool falls back to snippet from ir_find_documents. If snippet is also absent, tool returns a descriptive error."),
        ("Password-protected or paywalled documents", "None",
         "Access blocked. Tool returns a descriptive error."),
    ]
    for vals in fmt_rows:
        row = tbl_fmt.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_heading("JSON Config Location", level=3)
    doc.add_paragraph("sajhamcpserver/config/tools/ir_extract_content.json")

    doc.add_heading("Implementation Class", level=3)
    doc.add_paragraph("File:  sajhamcpserver/sajha/tools/impl/tavily_ir_tool.py")
    doc.add_paragraph("Class: IRExtractContentTool")

    # ════════════════════════════════════════════════════════════════════════
    # 5. Context Handling
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Context Handling", level=1)

    doc.add_heading("5.1 llm_extract() Compression \u2014 Primary Mechanism", level=2)
    doc.add_paragraph(
        "The primary context management mechanism is the secondary LLM extraction call in "
        "ir_extract_content. Raw content from Tavily extract (potentially large, unstructured HTML "
        "or PDF text) is passed to llm_extract() with a targeted extraction prompt built from the "
        "user query. The extraction call returns compact structured JSON containing only the "
        "information relevant to the query \u2014 this is the same pattern used by all EDGAR tools "
        "in edgar_tavily_tools.py."
    )
    doc.add_paragraph(
        "This is not optional. Do not pass raw Tavily extract output directly to the main LLM. "
        "Without llm_extract(), the main LLM receives unstructured HTML/text which degrades response "
        "quality and wastes context. The llm_extract() step is always required in ir_extract_content."
    )

    doc.add_heading("5.2 Fallback Strategy", level=2)
    doc.add_paragraph(
        "ir_extract_content uses a two-tier fallback when the primary extract call fails or returns "
        "insufficient content:"
    )
    tbl_fb = doc.add_table(rows=1, cols=3)
    tbl_fb.style = "Table Grid"
    make_header_row(tbl_fb, "Tier", "Method", "Trigger Condition")
    fb_rows = [
        ("1 \u2014 Primary", "tavily_extract(url) \u2192 llm_extract(content, prompt)",
         "Default path. Used when Tavily extract returns >200 chars of content."),
        ("2 \u2014 Fallback",
         "llm_extract(snippet, prompt) using snippet from ir_find_documents",
         "Triggered when Tavily extract returns <200 chars (scanned PDF, JS-rendered page, timeout). "
         "Caller must pass snippet= parameter from the ir_find_documents result for this URL."),
    ]
    for vals in fb_rows:
        row = tbl_fb.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph(
        "If both tiers return insufficient content, the tool returns a structured error response "
        "with a clear message: \"Could not extract content from {url}. The document may be a "
        "scanned PDF or require authentication. Try accessing the document directly.\""
    )

    # CHANGE 3: 5.3 Context Handling — fix the size guardrail paragraph
    doc.add_heading("5.3 Alignment with EDGAR Implementation", level=2)
    doc.add_paragraph(
        "The EDGAR tool stack (edgar_tavily_client.py, edgar_tavily_tools.py) has evolved to use a "
        "three-tier extraction strategy: (1) stream_sec_section() for large 10-K/10-Q HTML documents "
        "fetched directly from SEC Archives, (2) tavily_extract() for smaller or structured "
        "documents, (3) tavily_search() as last-resort fallback with _validate_sources() gate. EDGAR "
        "also uses direct_sec_json() to bypass Tavily entirely for SEC JSON API endpoints "
        "(submissions, XBRL concept data) where the URL is deterministic and the response is "
        "structured JSON."
    )
    doc.add_paragraph(
        "IR tools do not need stream_sec_section() or direct_sec_json() because IR documents are not "
        "SEC Archives HTML filings and their URLs are not deterministic until discovered via "
        "ir_find_page/ir_find_documents. The two-tier IR pattern (extract \u2192 snippet fallback) "
        "is the appropriate simplification for the IR use case. The fix_tavily_json() and "
        "llm_extract() utilities from edgar_tavily_client.py are shared and must be used by the IR "
        "implementation."
    )
    doc.add_paragraph(
        "As a safety net, ir_extract_content checks the byte length of the raw Tavily extract "
        "content BEFORE passing to llm_extract(). The llm_extract() function already limits input to "
        "6,000 chars and returns compact JSON \u2014 its output will never exceed 100KB. The "
        "byte-length check must guard the raw extracted content, not the structured output."
    )

    tbl_sz = doc.add_table(rows=1, cols=2)
    tbl_sz.style = "Table Grid"
    make_header_row(tbl_sz, "Response Size", "Action")
    sz_rows = [
        ("< 50 KB", "Pass through \u2014 return content to LLM as-is"),
        ("50 KB \u2013 100 KB",
         "Return content with a flag: content_truncated=true and a note in the response that the "
         "document is large and results may be partial"),
        ("> 100 KB",
         'Reject \u2014 return an error: "Extracted content too large ({size}KB). Refine your query '
         'to target a specific section or metric."'),
    ]
    for vals in sz_rows:
        row = tbl_sz.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph(
        "The 100KB rejection threshold is consistent with the universal size guardrail defined in "
        "the EDGAR Architecture Plan and applies to BaseMCPTool across all tools."
    )

    # ════════════════════════════════════════════════════════════════════════
    # 6. Files to Create
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Files to Create", level=1)

    doc.add_heading("6.1 New Tool Implementation", level=2)
    tbl_new1 = doc.add_table(rows=1, cols=2)
    tbl_new1.style = "Table Grid"
    make_header_row(tbl_new1, "File Path", "Description")
    row = tbl_new1.add_row()
    row.cells[0].text = "sajhamcpserver/sajha/tools/impl/tavily_ir_tool.py"
    row.cells[1].text = (
        "Single file containing all 3 tool classes: IRFindPageTool, IRFindDocumentsTool, "
        "IRExtractContentTool. Follows the same pattern as tavily_yahoo_finance_tool.py."
    )
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

    doc.add_heading("6.2 New JSON Tool Configs", level=2)
    tbl_new2 = doc.add_table(rows=1, cols=3)
    tbl_new2.style = "Table Grid"
    make_header_row(tbl_new2, "File Path", "Tool Name", "Class Reference")
    new2_rows = [
        ("config/tools/ir_find_page.json", "ir_find_page",
         "sajha.tools.impl.tavily_ir_tool.IRFindPageTool"),
        ("config/tools/ir_find_documents.json", "ir_find_documents",
         "sajha.tools.impl.tavily_ir_tool.IRFindDocumentsTool"),
        ("config/tools/ir_extract_content.json", "ir_extract_content",
         "sajha.tools.impl.tavily_ir_tool.IRExtractContentTool"),
    ]
    for vals in new2_rows:
        row = tbl_new2.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_heading("6.3 Implementation Pattern Reference", level=2)
    doc.add_paragraph(
        "tavily_ir_tool.py must follow the same structural pattern as edgar_tavily_tools.py \u2014 "
        "NOT tavily_yahoo_finance_tool.py. The EDGAR pattern is the correct reference because "
        "ir_extract_content uses extract (not search) and requires llm_extract(). Key elements:"
    )
    impl_bullets = [
        "Import tavily_extract, tavily_search, llm_extract, fix_tavily_json from "
        "edgar_tavily_client.py \u2014 do NOT re-implement these functions.",
        "Each tool class extends BaseMCPTool with get_input_schema(), get_output_schema(), and "
        "execute() methods.",
        'ir_find_page: calls tavily_search() with search_depth="basic", max_results=3, '
        "include_domains not set (open search).",
        "ir_find_documents: calls tavily_search() building the query from the document_type mapping "
        "table in Section 4.2, scoped to ir_domain + sec.gov.",
        "ir_extract_content: calls tavily_extract([url]) \u2192 fix_tavily_json() if sec.gov URL "
        "\u2192 llm_extract(content, prompt). Falls back to snippet parameter if extract returns "
        "<200 chars.",
        "Size check in ir_extract_content applied to raw Tavily extract content BEFORE passing to "
        "llm_extract(). The llm_extract() function already limits input to 6,000 chars and returns "
        "compact JSON \u2014 its output will never exceed 100KB. The byte-length check must guard "
        "the raw extracted content, not the structured output: "
        "if len(raw_content.encode()) > 100_000, raise ValueError with descriptive message.",
        "All tool responses include _source field set to the relevant URL.",
    ]
    for b in impl_bullets:
        doc.add_paragraph(b, style="List Bullet")

    # ════════════════════════════════════════════════════════════════════════
    # 7. Impact Summary
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("7. Impact Summary", level=1)

    tbl_imp = doc.add_table(rows=1, cols=3)
    tbl_imp.style = "Table Grid"
    make_header_row(tbl_imp, "Metric", "Before", "After")
    imp_rows = [
        ("IR tools", "7 tools", "3 tools"),
        ("Company coverage", "10 hardcoded companies", "Any publicly traded company"),
        ("Document content extraction", "Not available \u2014 URLs only",
         "Full text extraction via Tavily extract"),
        ("Scraper maintenance", "Required \u2014 breaks on site redesigns",
         "None \u2014 search-based, no HTML parsing"),
        ("Module size", "20 files (sajha/ir/) + 555-line impl", "1 impl file (~200 lines)"),
        ("Context safety", "None \u2014 no size checks",
         "Query-focused extraction + 100KB guardrail"),
        ("Dependencies removed", "\u2014",
         "requests/BeautifulSoup scraping stack in sajha/ir/"),
        # CHANGE 9: new row
        ("SEC Archives extraction", "Not supported",
         "stream_sec_section() for large HTML filings"),
    ]
    for vals in imp_rows:
        row = tbl_imp.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # 7.1 Risks
    doc.add_heading("7.1 Risks", level=2)
    tbl_risk = doc.add_table(rows=1, cols=3)
    tbl_risk.style = "Table Grid"
    make_header_row(tbl_risk, "Risk", "Likelihood", "Mitigation")
    risk_rows = [
        ("Scanned PDF documents return no content", "Low-Medium",
         "Tool returns clear error. Agent can suggest user access document directly."),
        ("Tavily search returns aggregator sites instead of official IR page", "Low",
         "ir_find_page filters results \u2014 skip SEC.gov, Yahoo Finance, Morningstar from "
         "ir_find_page results. ir_find_documents scopes search to ir_domain OR sec.gov explicitly."),
        # CHANGE 8: updated rate limits row
        ("Tavily rate limits under high query volume", "Low",
         "Add single-retry with 2-second backoff in tavily_search() and tavily_extract() for HTTP "
         "429 responses. Currently the functions raise ValueError('Rate limit exceeded') with no "
         "retry."),
        ("Very large IR pages with many documents \u2014 ir_find_documents returns incomplete list",
         "Medium",
         "Acceptable trade-off. Search ranking returns most relevant results. For exhaustive "
         "listing, SEC EDGAR submissions API (separate tool, future work) is the right approach."),
    ]
    for vals in risk_rows:
        row = tbl_risk.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # ════════════════════════════════════════════════════════════════════════
    # 8. Acceptance Criteria
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("8. Acceptance Criteria", level=1)
    tbl_ac = doc.add_table(rows=1, cols=3)
    tbl_ac.style = "Table Grid"
    make_header_row(tbl_ac, "#", "Criterion", "Pass Condition")
    ac_rows = [
        ("AC-01", "ir_find_page works for any major public company",
         "Returns a valid IR URL for at least 10 test companies spanning different IR platforms "
         "(Q4, Workiva, custom)"),
        ("AC-02", "ir_find_documents returns relevant documents",
         "For AAPL annual_report 2024, returns at least one result with a sec.gov or apple.com URL"),
        ("AC-03", "ir_extract_content reads earnings press release",
         'Extracts revenue and EPS figures from a known earnings press release URL with '
         'query="revenue EPS"'),
        ("AC-04", "Size guardrail triggers correctly",
         "Extraction of a known large document (full 10-K PDF) returns a rejection error with "
         "descriptive message"),
        ("AC-05", "No existing IR tool names resolve",
         "Calls to ir_get_annual_reports, ir_get_presentations, ir_list_supported_companies, etc. "
         "return tool-not-found errors after removal"),
        ("AC-06", "IRIS CCR tools unaffected",
         "All 9 iris_* tools continue to function after the IR removal"),
        ("AC-07", "Small company coverage",
         "ir_find_page and ir_find_documents work for a mid-cap company not previously in the "
         "hardcoded database"),
    ]
    for vals in ac_rows:
        row = tbl_ac.add_row()
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # Save
    out_path = f"{OUTPUT_DIR}/Sajha_IR_Tavily_Replacement_ERD.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


# ─────────────────────────────────────────────────────────────────────────────
# FILE 2 — Tavily_SEC_MDA_Direct_Fetch_Note.docx
# ─────────────────────────────────────────────────────────────────────────────

def build_mda_note():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Title ───────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph(
        "Why Tavily Cannot Extract SEC MD&A \u2014 Direct Streaming as the Only Reliable Approach"
    )
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.runs[0].bold = True
    title_p.runs[0].font.size = Pt(16)
    title_p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub_p = doc.add_paragraph(
        "Technical Note \u2014 RiskGPT Engineering  |  March 2026"
    )
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.runs[0].font.size = Pt(10)
    sub_p.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    sub_p.runs[0].italic = True

    doc.add_paragraph()

    # ── Section: The Problem ─────────────────────────────────────────────────
    doc.add_heading("The Problem", level=1)
    doc.add_paragraph(
        "Tavily\u2019s /extract endpoint is designed for typical web pages (10\u2013500 KB). "
        "SEC EDGAR 10-K annual report filings for large financial institutions routinely exceed "
        "10 MB of HTML. Bank of America\u2019s FY2025 10-K (bac-20251231.htm, filed 2026-02-25) "
        "is 12,819,060 bytes \u2014 12.2 MB. When Tavily /extract is called on a URL this large, "
        "it either times out (the 20-second timeout is reached before the document is parsed) or "
        "returns cached/partial content from its web index that may belong to entirely different "
        "companies. In testing, a query for \u201cBAC Q4 2025 MD&A\u201d returned content "
        "attributed to Tracking Inc., Invesco, and Vital Farms \u2014 three unrelated companies. "
        "The agent synthesised a plausible-sounding response from this garbage data without "
        "flagging any issues. This is the most dangerous failure mode: silent wrong-company synthesis."
    )

    # ── Section: Why Tavily /search Is No Better ─────────────────────────────
    doc.add_heading("Why Tavily /search Is No Better", level=1)
    doc.add_paragraph(
        "Switching to Tavily /search with include_domains=['sec.gov'] does not fix the problem. "
        "Keyword search returns whichever SEC document best matches the search terms \u2014 not "
        "necessarily the one the user requested. In the BAC test, /search returned a Disney filing "
        "and a 2008 BAC filing (17 years stale) for a \u201cQ4 2025\u201d query. The keyword match "
        "on \u201cmanagement discussion analysis Bank of America\u201d retrieved documents where "
        "those terms appeared together \u2014 regardless of the filing year or company CIK."
    )

    # ── Section: Why the EDGAR Submissions API Alone Is Not Sufficient ────────
    doc.add_heading(
        "Why the EDGAR Submissions API Alone Is Not Sufficient", level=1
    )
    doc.add_paragraph(
        "The SEC submissions API (data.sec.gov/submissions/CIK{padded}.json) reliably identifies "
        "the correct filing URL. After implementing the submissions API lookup, we had the exact URL "
        "for the correct document: "
        "https://www.sec.gov/Archives/edgar/data/70858/000007085826000157/bac-20251231.htm. "
        "But passing this URL to tavily_extract returned the same wrong-company content as before. "
        "The issue is not URL resolution \u2014 it is document retrieval at scale."
    )

    # ── Section: Why Direct Streaming Is the Only Reliable Approach ───────────
    doc.add_heading(
        "Why Direct Streaming Is the Only Reliable Approach", level=1
    )
    doc.add_paragraph(
        "The SEC Archives are publicly accessible static HTML files. A direct urllib.request call "
        "with a proper User-Agent header (required by SEC) fetches the raw HTML bytes. The document "
        "is read in 64 KB chunks. The MD&A section (Item 7) is located by scanning for the section "
        "marker text \u2014 the first occurrence is in the Table of Contents, the second or third "
        "is the actual section body. Once the marker position is found, the next 120 KB of text is "
        "extracted, HTML tags are stripped, and clean plain text is returned."
    )

    doc.add_paragraph("Key facts about this approach:")
    facts = [
        "BAC\u2019s Item 7 starts at byte position ~1,810,578 (approximately 1.8 MB into the file)",
        "We read ~1.8 MB to reach the section, then extract 120 KB of content",
        "Total bytes transferred: ~2 MB vs 12.2 MB full document",
        "No Tavily quota consumed for SEC Archives HTML",
        "Validated against BAC FY2025 (10-K), AAPL FY2024 (10-K), JPM Q3 2025 (10-Q) \u2014 all pass",
    ]
    for f in facts:
        doc.add_paragraph(f, style="List Bullet")

    # ── Section: Implementation ──────────────────────────────────────────────
    doc.add_heading("Implementation", level=1)
    doc.add_paragraph(
        "The stream_sec_section(filing_url, section_marker, content_kb=120) function in "
        "edgar_tavily_client.py implements this approach. It is called as Tier 1 in the "
        "_extract_from_filing_url() shared helper used by all four EDGAR qualitative tools "
        "(edgar_extract_section, edgar_earnings_brief, edgar_segment_analysis, edgar_risk_summary). "
        "The function is also the required fallback in ir_extract_content when the document URL is "
        "from sec.gov/Archives."
    )

    out_path = f"{OUTPUT_DIR}/Tavily_SEC_MDA_Direct_Fetch_Note.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_erd()
    build_mda_note()
    print("Done.")
